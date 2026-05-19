import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    # Highlight with light gray background
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F2F2F2')
    p.paragraph_format.element.get_or_add_pPr().append(shading_elm)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    return p

def add_bullet(doc, title, text):
    p = doc.add_paragraph(style='List Bullet')
    if title:
        run = p.add_run(title)
        run.bold = True
    if text:
        p.add_run(text)

def add_highlight_box(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x00) # Dark green
    p.paragraph_format.left_indent = Inches(0.4)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E8F5E9') # Light green background
    p.paragraph_format.element.get_or_add_pPr().append(shading_elm)

doc = Document()

# Title
title = doc.add_heading('Cẩm nang Phỏng vấn: Database Specialist - Corporate Banking Migration', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Tài liệu hướng dẫn chi tiết dành cho vị trí Database Specialist, tập trung vào thiết kế, tối ưu hóa (Performance Tuning) và vận hành hệ thống cơ sở dữ liệu lớn (Oracle/PostgreSQL). Tài liệu bao gồm các Case Study thực chiến đã được tổng hợp.')

# CASE STUDY 1
doc.add_heading('1. Kịch Bản Thiết Kế ERD & Bảng Biểu (System Design)', level=1)
doc.add_paragraph('Phân hệ lưu trữ giao dịch và xử lý cảnh báo rửa tiền (AML) với 4 bảng cốt lõi:')
add_bullet(doc, 'CUSTOMER (1) — (N) ACCOUNT: ', 'Khách hàng có nhiều tài khoản.')
add_bullet(doc, 'ACCOUNT (1) — (N) TRANSACTION: ', 'Tài khoản phát sinh nhiều giao dịch.')
add_bullet(doc, 'TRANSACTION (1) — (0..N) AML_ALERT: ', 'Giao dịch sinh ra cảnh báo rủi ro.')

doc.add_heading('Thiết kế chi tiết cấu trúc (Oracle 19c DDL)', level=2)
sql_model = """-- 1. BẢNG DANH MỤC KHÁCH HÀNG
CREATE TABLE CUSTOMER (
    CUSTOMER_ID   NUMBER(19,0) NOT NULL,
    FULL_NAME     VARCHAR2(150) NOT NULL,
    ID_CARD       VARCHAR2(256) NOT NULL, -- Sẽ được Encrypt bằng DBMS_CRYPTO (hoặc pgcrypto)
    PHONE         VARCHAR2(256),          
    STATUS        VARCHAR2(20) DEFAULT 'ACTIVE',
    CREATED_AT    TIMESTAMP DEFAULT SYSTIMESTAMP,
    CONSTRAINT PK_CUSTOMER PRIMARY KEY (CUSTOMER_ID)
);

-- 2. BẢNG TÀI KHOẢN NGÂN HÀNG
CREATE TABLE ACCOUNT (
    ACCOUNT_NO    VARCHAR2(20) NOT NULL,
    CUSTOMER_ID   NUMBER(19,0) NOT NULL,
    ACCT_TYPE     VARCHAR2(10) NOT NULL,
    BALANCE       NUMBER(18,2) DEFAULT 0.00, -- Tuyệt đối không dùng FLOAT để tránh sai số
    CURRENCY      VARCHAR2(3) DEFAULT 'VND',
    STATUS        VARCHAR2(20) DEFAULT 'OPEN',
    CREATED_AT    TIMESTAMP DEFAULT SYSTIMESTAMP,
    CONSTRAINT PK_ACCOUNT PRIMARY KEY (ACCOUNT_NO),
    CONSTRAINT FK_ACCOUNT_CUSTOMER FOREIGN KEY (CUSTOMER_ID) REFERENCES CUSTOMER(CUSTOMER_ID)
);"""
add_code_block(doc, sql_model)

add_highlight_box(doc, "💡 Điểm Specialist cần nhấn mạnh:\n- Kiểu dữ liệu: Số tiền luôn dùng NUMBER(18,2) (hoặc DECIMAL trong Postgres), không bao giờ dùng FLOAT.\n- Bảo mật: PII (ID_CARD, PHONE) mã hóa AES256 từ vùng ODS (bằng DBMS_CRYPTO hoặc pgcrypto) theo luật An toàn thông tin.")

doc.add_heading('Thiết kế Lớp Ứng Dụng (Application Layer) Tương Tác CSDL', level=2)
add_bullet(doc, 'Connection Pooling: ', 'Ứng dụng kết nối qua HikariCP. Giới hạn maximumPoolSize hợp lý để tránh kẹt DB.')
add_bullet(doc, 'Caching: ', 'Bảng CUSTOMER ít biến động -> Cache trên Redis ở lớp ứng dụng. Bảng TRANSACTION biến động liên tục -> Không cache, đọc trực tiếp DB hoặc Replica.')
add_bullet(doc, 'Phân tách Đọc/Ghi (CQRS): ', 'Giao dịch ghi đi vào DB Master. Truy vấn lịch sử giao dịch đi vào DB Standby để giảm tải Master.')

# CASE STUDY 2
doc.add_heading('2. Kịch Bản Đánh Index & Partitioning (Bảng Tỷ Dòng)', level=1)
doc.add_paragraph('Đối với bảng TRANSACTION lên tới trăm triệu dòng mỗi tháng, quy hoạch vật lý là bắt buộc.')
sql_partition = """-- 3. BẢNG GIAO DỊCH (TỶ DÒNG) - Bỏ Physical Foreign Key để tối ưu INSERT
CREATE TABLE TRANSACTION (
    TXN_ID         NUMBER(19,0) NOT NULL,
    ACCOUNT_NO     VARCHAR2(20) NOT NULL,
    TXN_TYPE       VARCHAR2(10) NOT NULL,
    AMOUNT         NUMBER(18,2) NOT NULL,
    TXN_DATE       DATE NOT NULL,         -- Khóa phân vùng (Partition Key)
    DESCRIPTION    VARCHAR2(250),
    CHANNEL        VARCHAR2(20)
)
PARTITION BY RANGE (TXN_DATE) (
    PARTITION SUB_TXN_2026_M04 VALUES LESS THAN (TO_DATE('2026-05-01', 'YYYY-MM-DD')),
    PARTITION SUB_TXN_2026_M05 VALUES LESS THAN (TO_DATE('2026-06-01', 'YYYY-MM-DD')),
    PARTITION SUB_TXN_DEFAULT VALUES LESS THAN (MAXVALUE)
);

-- TẠO LOCAL INDEX
CREATE INDEX IX_TXN_ACCT_DATE ON TRANSACTION (ACCOUNT_NO, TXN_DATE) LOCAL;

-- 4. BẢNG CẢNH BÁO AML
CREATE TABLE AML_ALERT (
    ALERT_ID      NUMBER(19,0) NOT NULL,
    TXN_ID        NUMBER(19,0) NOT NULL,
    RULE_CODE     VARCHAR2(30) NOT NULL,
    RISK_LEVEL    VARCHAR2(10),
    ALERT_DATE    TIMESTAMP DEFAULT SYSTIMESTAMP,
    STATUS        VARCHAR2(20) DEFAULT 'NEW',
    CONSTRAINT PK_AML_ALERT PRIMARY KEY (ALERT_ID)
);"""
add_code_block(doc, sql_partition)

doc.add_heading('Chiến lược giải thích với Khách hàng:', level=2)
add_bullet(doc, 'Partition Pruning: ', 'Truy vấn theo tháng (TXN_DATE) sẽ chỉ quét đúng partition của tháng đó, tốc độ tăng gấp hàng trăm lần.')
add_bullet(doc, 'Data Purging an toàn: ', 'Khi cần dọn data 5 năm trước, thay vì dùng lệnh DELETE gây treo DB, chỉ cần dùng ALTER TABLE TRANSACTION DROP PARTITION (chạy mất 0.5s).')
add_bullet(doc, 'Local Index: ', 'Khách hàng thường tra cứu lịch sử của 1 tài khoản trong 1 tháng. Dùng Local Index để khi Drop partition, các index khác vẫn USABLE.')
add_bullet(doc, 'Bỏ Physical FK: ', 'Đề xuất không tạo Foreign Key trỏ về bảng ACCOUNT để tối đa tốc độ INSERT realtime. Việc check tính toàn vẹn đưa lên tầng Application hoặc ETL.')

# CASE STUDY 3
doc.add_heading('3. Công Việc Hằng Ngày (Daily Tasks & Troubleshooting)', level=1)
doc.add_paragraph('Trả lời mạch lạc theo quy trình 3 giai đoạn của DBA/Specialist:')
doc.add_heading('Ca sáng (08h00 - 09h30): Health Check', level=2)
add_bullet(doc, '', 'Kiểm tra Dashboard Grafana / Dynatrace: CPU, RAM, IOPS, Active Sessions, Buffer Hit Ratio (> 95%).')
add_bullet(doc, '', 'Kiểm tra trạng thái đồng bộ HA: Oracle Data Guard (Apply Lag) hoặc PostgreSQL Replication Slot xem có bị lag không.')
add_bullet(doc, '', 'Check các Batch Job ETL đêm qua xem có job nào FAIL hoặc Long-running không.')

doc.add_heading('Ca ngày (09h30 - 16h30): Performance Tuning', level=2)
add_bullet(doc, 'Bắt Slow Query: ', 'Dùng AWR Report (Oracle) hoặc pg_stat_statements (PostgreSQL).')
add_bullet(doc, '', 'Dùng EXPLAIN PLAN / EXPLAIN ANALYZE để xem câu lệnh bị Full Table Scan hay Optimizer tính sai Cardinality.')
add_bullet(doc, '', 'Phối hợp Dev review SQL trước khi deploy lên UAT/PROD.')

doc.add_heading('Cuối ngày / Cuối tuần: Bảo trì (Maintenance)', level=2)
add_bullet(doc, '', 'Theo dõi dung lượng Tablespace (Capacity Planning).')
add_bullet(doc, '', 'Chạy script thu gom rác (VACUUM trên PostgreSQL) hoặc cập nhật thống kê (DBMS_STATS) vào ban đêm.')

add_highlight_box(doc, "🔥 Vũ khí bí mật - Quy trình xử lý sự cố (Treo DB):\n1. Tìm Root Cause: Không đoán mò. Check Wait Events bằng V$SESSION_WAIT / ASH (Oracle) hoặc pg_stat_activity / pg_locks (Postgres).\n2. Workaround khẩn cấp: Xin phê duyệt KILL SESSION gây nghẽn luồng giao dịch.\n3. Permanent Fix: Mang câu SQL về Test mổ xẻ Plan, thêm Index hoặc chỉnh Partition để dứt điểm.")

# CASE STUDY 4
doc.add_heading('4. Kịch Bản Xử Lý Xung Đột Giữa ETL Batch & OLTP', level=1)
doc.add_paragraph('Tình huống thường gặp: Job ODI/Talend chạy load bảng dữ liệu nghiệp vụ vào ban đêm kéo dài lấn sang giờ hành chính, gây Row Lock Contention hoặc làm treo Core Banking.')
doc.add_heading('Góc nhìn của Data Engineer (Thiết kế luồng ETL)', level=2)
add_bullet(doc, 'Tránh Update/Merge trực tiếp khối lượng lớn: ', 'Sử dụng Partition Exchange thay vì Update trực tiếp.')
add_bullet(doc, 'Quy trình: ', 'Load dữ liệu vào một bảng tạm (Staging table) có cấu trúc y hệt. Xây dựng index đầy đủ trên bảng tạm. Sau đó dùng lệnh ALTER TABLE ... EXCHANGE PARTITION. Lệnh này chỉ hoán đổi metadata, diễn ra trong vòng chưa tới 1 giây.')
add_bullet(doc, 'Giới hạn DOP (Degree of Parallelism): ', 'Đảm bảo cấu hình Resource Manager không cho phép session ETL ngốn quá 40% CPU của hệ thống.')

doc.add_heading('Góc nhìn của Database Administrator (DBA Monitoring)', level=2)
sql_block = """-- DBA phát hiện job ODI đang block user Online
SELECT s1.username || '@' || s1.machine AS blocker_session,
       s2.username || '@' || s2.machine AS blocked_session,
       s2.wait_class, s2.seconds_in_wait
FROM gv$lock l1, gv$session s1, gv$lock l2, gv$session s2
WHERE s1.sid = l1.sid AND s2.sid = l2.sid
  AND l1.BLOCK = 1 AND l2.request > 0
  AND l1.id1 = l2.id1 AND l1.id2 = l2.id2;"""
add_code_block(doc, sql_block)

add_highlight_box(doc, "🔥 Vũ khí bí mật khi phỏng vấn:\nĐừng bao giờ nói \"Em sẽ Kill Session của job ODI\". Hãy trả lời:\n1. Tạm dừng Agent của ODI/Talend (Graceful shutdown).\n2. Phối hợp với team nghiệp vụ dời lịch chạy batch_window sang khung giờ thấp điểm.\n3. Đề xuất quy hoạch lại IKM trong ODI: Sử dụng IKM Oracle Control Append kết hợp nạp theo từng Chunk nhỏ.")

# CASE STUDY 5
doc.add_heading('5. Kịch Bản Data Architecture & Migration (Corporate Banking)', level=1)
doc.add_paragraph('Hệ thống Corporate Banking (Khách hàng Doanh nghiệp) với hàng trăm Entity phức tạp yêu cầu cách tiếp cận Migration an toàn, Multi-Phase.')

doc.add_heading('Chiến lược Migration: Strangler Fig Pattern & Parallel Run', level=2)
add_bullet(doc, 'Strangler Fig Pattern: ', 'Không đập bỏ toàn bộ Core Legacy cùng lúc. Chia nhỏ từng domain. Lớp API Gateway sẽ làm nhiệm vụ Router.')
add_bullet(doc, 'Change Data Capture (CDC): ', 'Dùng Kafka Connect + Debezium (hoặc Oracle GoldenGate) để đồng bộ (Sync) dữ liệu 2 chiều real-time giữa DB cũ và DB mới. Mục đích: Fallback (quay xe) về hệ thống cũ ngay lập tức mà không bị thất thoát giao dịch.')
add_bullet(doc, 'Parallel Run & Reconciliation (Đối soát): ', 'Chạy ngầm cả hai hệ thống trong 1-3 tháng. Thiết kế các Batch Job (ETL) chạy cuối ngày để đối chiếu Tổng tài sản, Số dư. Nếu sai số = 0 trong 30 ngày liên tục mới được Sign-off.')

doc.add_heading('Thiết kế Kiến trúc Dữ liệu mới (Microservices Data Architecture)', level=2)
add_bullet(doc, 'Database-per-service (Tách DB): ', 'Phá vỡ DB Monolithic khổng lồ. Giải quyết triệt để vấn đề "Một bảng chết kéo theo toàn hệ thống chết".')
add_bullet(doc, 'CQRS & Tách biệt OLTP/OLAP: ', 'Database Core chỉ xử lý Transaction. Dữ liệu được stream (qua Kafka) về Data Lake/Warehouse để chạy báo cáo cuối ngày.')
add_bullet(doc, 'Event-Driven & SAGA Pattern: ', 'Trong kiến trúc mới không thể JOIN 2 bảng ở 2 DB khác nhau. Phải áp dụng SAGA Pattern để xử lý transaction phân tán.')

add_highlight_box(doc, "💡 Điểm Data Architect cần \"Đóng đinh\" hội đồng:\n- \"Thưa anh/chị, Migration không phải là bài toán Copy SQL Data, mà cốt lõi là bài toán Đối soát (Reconciliation) và Fallback Plan (Kế hoạch lùi).\"\n- \"Xây DB Microservices không khó ở việc tạo bảng, mà khó ở việc quản trị Master Data (MDM) để tránh dữ liệu rác, và đảm bảo Eventual Consistency (Nhất quán cuối) khi chia tách dữ liệu.\"")

doc.save('Database_Specialist_Interview_Guide_VN_v2.docx')
